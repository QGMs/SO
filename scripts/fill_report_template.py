#!/usr/bin/env python3
"""Fill an existing DOCX report template while preserving its cover and index pages."""

from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


HEADING_STYLE_BY_LEVEL = {
    1: "Heading 1",
    2: "Heading 2",
    3: "Heading 3",
    4: "Heading 3",
    5: "Heading 3",
    6: "Heading 3",
}


def add_inline_runs(paragraph, text: str) -> None:
    token_re = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0

    for match in token_re.finditer(text):
        start, end = match.span()
        if start > pos:
            paragraph.add_run(text[pos:start])

        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            paragraph.add_run(token)
        pos = end

    if pos < len(text):
        paragraph.add_run(text[pos:])


def split_cover_and_body(markdown_text: str) -> tuple[List[str], List[str]]:
    lines = markdown_text.splitlines()
    for idx, line in enumerate(lines):
        if line.strip() == "---":
            return lines[:idx], lines[idx + 1 :]
    return lines, []


def parse_table_row(line: str) -> List[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_table_delimiter_row(cells: List[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", cell.strip()) is not None for cell in cells)


def apply_table_font(table) -> None:
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_after = Pt(3)
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10.5)
                    if row_index == 0:
                        run.bold = True


def find_anchor_element(doc: Document, anchor_text: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == anchor_text:
            return paragraph._element
    raise ValueError(f"Anchor paragraph not found: {anchor_text}")


def clear_body_from_anchor(doc: Document, anchor_text: str) -> None:
    body = doc._element.body
    anchor = find_anchor_element(doc, anchor_text)
    remove = False
    for child in list(body):
        if child is anchor:
            remove = True
        if remove and child.tag != body.sectPr.tag:
            body.remove(child)


def add_normal_paragraph(doc: Document, text: str, indent_cm: float = 0.0) -> None:
    paragraph = doc.add_paragraph(style="normal")
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.left_indent = Cm(indent_cm)
    add_inline_runs(paragraph, text)


def add_bullet_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="normal")
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Cm(0.55)
    paragraph.paragraph_format.first_line_indent = Cm(-0.35)
    add_inline_runs(paragraph, f"- {text}")


def add_number_paragraph(doc: Document, text: str, number: str) -> None:
    paragraph = doc.add_paragraph(style="normal")
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Cm(0.55)
    paragraph.paragraph_format.first_line_indent = Cm(-0.35)
    add_inline_runs(paragraph, f"{number}. {text}")


def add_code_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="normal")
    paragraph.paragraph_format.left_indent = Cm(0.7)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)


def add_heading(doc: Document, markdown_level: int, text: str) -> None:
    effective_level = max(1, markdown_level - 1)
    style = HEADING_STYLE_BY_LEVEL.get(effective_level, "Heading 3")
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_inline_runs(paragraph, text)


def add_markdown_body(doc: Document, body_lines: List[str]) -> None:
    i = 0
    while i < len(body_lines):
        line = body_lines[i].rstrip("\n")
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph(style="normal")
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph(style="normal")
            i += 1
            continue

        if stripped.startswith("```"):
            i += 1
            while i < len(body_lines) and not body_lines[i].strip().startswith("```"):
                add_code_paragraph(doc, body_lines[i].rstrip("\n"))
                i += 1
            if i < len(body_lines):
                i += 1
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            add_heading(doc, len(heading_match.group(1)), heading_match.group(2).strip())
            i += 1
            continue

        if stripped.startswith("|"):
            raw_table_lines: List[str] = []
            while i < len(body_lines) and body_lines[i].strip().startswith("|"):
                raw_table_lines.append(body_lines[i].strip())
                i += 1

            parsed_rows: List[List[str]] = []
            for raw in raw_table_lines:
                cells = parse_table_row(raw)
                if cells and not is_table_delimiter_row(cells):
                    parsed_rows.append(cells)

            if parsed_rows:
                column_count = max(len(row) for row in parsed_rows)
                table = doc.add_table(rows=len(parsed_rows), cols=column_count)
                try:
                    table.style = "Table Grid"
                except KeyError:
                    pass
                for row_index, row_cells in enumerate(parsed_rows):
                    for col_index, value in enumerate(row_cells):
                        cell_paragraph = table.cell(row_index, col_index).paragraphs[0]
                        cell_paragraph.text = ""
                        add_inline_runs(cell_paragraph, value)
                apply_table_font(table)
            continue

        bullet_match = re.match(r"^-\s+(.*)$", stripped)
        if bullet_match:
            add_bullet_paragraph(doc, bullet_match.group(1).strip())
            i += 1
            continue

        number_match = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if number_match:
            add_number_paragraph(doc, number_match.group(2).strip(), number_match.group(1))
            i += 1
            continue

        add_normal_paragraph(doc, stripped)
        i += 1


def generate_from_template(template_path: Path, markdown_path: Path, output_path: Path, anchor_text: str) -> None:
    markdown = markdown_path.read_text(encoding="utf-8")
    _, body_lines = split_cover_and_body(markdown)

    doc = Document(str(template_path))
    clear_body_from_anchor(doc, anchor_text)
    add_markdown_body(doc, body_lines)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("template")
    parser.add_argument("markdown")
    parser.add_argument("output")
    parser.add_argument("--anchor-text", default="Resumo")
    args = parser.parse_args()

    generate_from_template(
        template_path=Path(args.template),
        markdown_path=Path(args.markdown),
        output_path=Path(args.output),
        anchor_text=args.anchor_text,
    )
    print(f"Generated: {Path(args.output).resolve()}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

