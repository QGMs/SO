#!/usr/bin/env python3
"""
Convert a Markdown report to DOCX with a clean academic layout.

Features:
- A4 layout with compact margins.
- Calibri 11pt as base style.
- 1-page cover (split at first --- line).
- Left-aligned headings, justified body text.
- Support for headings, bullets, numbered lists, code fences, and tables.
- No header/footer content.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path
from typing import List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


HEADING_STYLE_BY_LEVEL = {
    1: "Heading 1",
    2: "Heading 2",
    3: "Heading 3",
    4: "Heading 3",
    5: "Heading 3",
    6: "Heading 3",
}


def configure_page(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)

        section.header_distance = Cm(0.0)
        section.footer_distance = Cm(0.0)
        for paragraph in section.header.paragraphs:
            paragraph.text = ""
        for paragraph in section.footer.paragraphs:
            paragraph.text = ""


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(8)

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(11)
    h3.font.bold = True
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)


def split_cover_and_body(markdown_text: str) -> tuple[List[str], List[str]]:
    lines = markdown_text.splitlines()
    for idx, line in enumerate(lines):
        if line.strip() == "---":
            return lines[:idx], lines[idx + 1 :]
    return lines, []


def add_inline_runs(paragraph, text: str) -> None:
    """
    Parse simple inline markdown:
    - **bold**
    - `code`
    """
    token_re = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0

    for match in token_re.finditer(text):
        start, end = match.span()
        if start > pos:
            paragraph.add_run(text[pos:start])

        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            paragraph.add_run(token)
        pos = end

    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_cover(doc: Document, cover_lines: List[str]) -> None:
    title_text = ""
    other_lines: List[str] = []

    for line in cover_lines:
        stripped = line.strip()
        if stripped.startswith("# "):
            title_text = stripped[2:].strip()
        elif stripped:
            other_lines.append(stripped)

    # Give visual room at top of cover.
    spacer = doc.add_paragraph("")
    spacer.paragraph_format.space_before = Pt(70)

    if title_text:
        p_title = doc.add_paragraph("", style="Heading 1")
        p_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_inline_runs(p_title, title_text)

    doc.add_paragraph("")

    for line in other_lines:
        paragraph = doc.add_paragraph("")
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(8)
        add_inline_runs(paragraph, line)

    doc.add_page_break()


def is_table_delimiter_row(cells: List[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", cell.strip()) is not None for cell in cells)


def parse_table_row(line: str) -> List[str]:
    raw_cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return raw_cells


def apply_table_font(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_after = Pt(3)
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10.5)


def add_markdown_body(doc: Document, body_lines: List[str]) -> None:
    i = 0
    while i < len(body_lines):
        line = body_lines[i].rstrip("\n")
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph("")
            i += 1
            continue

        # Code fence block.
        if stripped.startswith("```"):
            i += 1
            while i < len(body_lines) and not body_lines[i].strip().startswith("```"):
                code_paragraph = doc.add_paragraph("")
                code_paragraph.paragraph_format.left_indent = Cm(0.6)
                code_paragraph.paragraph_format.space_after = Pt(2)
                run = code_paragraph.add_run(body_lines[i].rstrip("\n"))
                run.font.name = "Consolas"
                run.font.size = Pt(10)
                i += 1
            if i < len(body_lines):
                i += 1
            continue

        # Heading.
        match_heading = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if match_heading:
            level = len(match_heading.group(1))
            text = match_heading.group(2).strip()
            style = HEADING_STYLE_BY_LEVEL.get(level, "Heading 3")
            paragraph = doc.add_paragraph("", style=style)
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline_runs(paragraph, text)
            i += 1
            continue

        # Table block.
        if stripped.startswith("|"):
            raw_table_lines: List[str] = []
            while i < len(body_lines) and body_lines[i].strip().startswith("|"):
                raw_table_lines.append(body_lines[i].strip())
                i += 1

            parsed_rows: List[List[str]] = []
            for raw in raw_table_lines:
                cells = parse_table_row(raw)
                if not cells:
                    continue
                if is_table_delimiter_row(cells):
                    continue
                parsed_rows.append(cells)

            if parsed_rows:
                col_count = max(len(r) for r in parsed_rows)
                table = doc.add_table(rows=len(parsed_rows), cols=col_count)
                table.style = "Table Grid"
                for r_idx, row_cells in enumerate(parsed_rows):
                    for c_idx, value in enumerate(row_cells):
                        cell_p = table.cell(r_idx, c_idx).paragraphs[0]
                        cell_p.text = ""
                        add_inline_runs(cell_p, value)
                apply_table_font(table)
            continue

        # Bullet list.
        match_bullet = re.match(r"^-\s+(.*)$", stripped)
        if match_bullet:
            paragraph = doc.add_paragraph("", style="List Bullet")
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline_runs(paragraph, match_bullet.group(1).strip())
            i += 1
            continue

        # Numbered list.
        match_number = re.match(r"^\d+\.\s+(.*)$", stripped)
        if match_number:
            paragraph = doc.add_paragraph("", style="List Number")
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline_runs(paragraph, match_number.group(1).strip())
            i += 1
            continue

        # Default paragraph.
        paragraph = doc.add_paragraph("")
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_inline_runs(paragraph, stripped)
        i += 1


def convert_markdown_to_docx(input_path: Path, output_path: Path) -> None:
    markdown = input_path.read_text(encoding="utf-8")

    doc = Document()
    configure_page(doc)
    configure_styles(doc)

    cover_lines, body_lines = split_cover_and_body(markdown)
    add_cover(doc, cover_lines)
    add_markdown_body(doc, body_lines)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main() -> int:
    input_arg = sys.argv[1] if len(sys.argv) > 1 else "relatorio.md"
    output_arg = sys.argv[2] if len(sys.argv) > 2 else "relatorio.docx"

    input_path = Path(input_arg)
    output_path = Path(output_arg)

    if not input_path.exists():
        print(f"Input not found: {input_path}", file=sys.stderr)
        return 1

    convert_markdown_to_docx(input_path, output_path)
    print(f"Generated: {output_path.resolve()}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
